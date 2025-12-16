import os
import asyncio
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from app.services.llm_engine import get_llm
from app.services.risk_checker import RiskCheckerService
from app.schemas.contract_schema import ContractDraftRequest, RiskAnalysisRequest

class DrafterService:
    def __init__(self):
        # Dùng model tốt (GPT-4o/Gemini Pro) để viết văn bản hay
        self.llm = get_llm(streaming=False, temperature=0.5)
        self.risk_checker = RiskCheckerService()

    async def draft_contract(self, data: ContractDraftRequest) -> dict:
        # 1. Soạn thảo văn bản
        prompt = ChatPromptTemplate.from_template("""
        <|im_start|>system
        Bạn là Luật sư cao cấp tại Việt Nam. Nhiệm vụ: Soạn thảo văn bản pháp lý chuyên nghiệp.
        
        YÊU CẦU CẤU TRÚC VĂN BẢN (BẮT BUỘC):
        1. Quốc hiệu: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM (Dòng 1)
        2. Tiêu ngữ: Độc lập - Tự do - Hạnh phúc (Dòng 2)
        3. Tên văn bản: VIẾT HOA TOÀN BỘ (VD: HỢP ĐỒNG MUA BÁN)
        4. Căn cứ pháp lý: Căn cứ Bộ luật Dân sự 2015...
        5. Nội dung: Chia thành "Điều 1:", "Điều 2:" rõ ràng.
        6. Phần ký tên: Bên A và Bên B ở cuối.

        Ngôn ngữ: Trang trọng, chặt chẽ, bảo vệ quyền lợi hợp pháp.
        <|im_end|>
        
        <|im_start|>user
        Yêu cầu soạn thảo:
        - Loại văn bản: {contract_type}
        - Bên A: {party_a}
        - Bên B: {party_b}
        - Chi tiết/Yêu cầu đặc biệt: {key_terms}
        <|im_end|>
        
        <|im_start|>assistant
        """)
        
        chain = prompt | self.llm | StrOutputParser()
        raw_content = await chain.ainvoke({
            "contract_type": data.contract_type,
            "party_a": data.party_a,
            "party_b": data.party_b,
            "key_terms": str(data.key_terms)
        })

        # 2. Chạy Risk Checker song song hoặc tuần tự
        # (Lưu ý: Risk checker cần input là RiskAnalysisRequest)
        risk_data_input = RiskAnalysisRequest(contract_type=data.contract_type, content=raw_content)
        risk_result = await self.risk_checker.analyze_document(risk_data_input)

        # 3. Xử lý kết quả Risk (Hỗ trợ cả Dict và Pydantic Model)
        if isinstance(risk_result, dict):
            score = risk_result.get("overall_score", 0)
            risks = risk_result.get("risks", [])
        else:
            # Nếu là Pydantic model
            score = getattr(risk_result, "overall_score", 0)
            risks = getattr(risk_result, "risks", [])

        # Tạo tóm tắt rủi ro
        high_risks = []
        for r in risks:
            # Xử lý linh hoạt dù r là dict hay object
            severity = r.get('severity') if isinstance(r, dict) else getattr(r, 'severity', '')
            issue = r.get('issue') if isinstance(r, dict) else getattr(r, 'issue', '')
            
            if str(severity).lower() == 'high':
                high_risks.append(issue)

        risk_summary = f"Điểm an toàn pháp lý: {score}/100\n"
        if high_risks:
            risk_summary += "⚠️ CÁC ĐIỂM CẦN LƯU Ý:\n- " + "\n- ".join(high_risks[:3]) # Lấy top 3 lỗi
            if len(high_risks) > 3: risk_summary += f"\n... và {len(high_risks)-3} vấn đề khác."
        else:
            risk_summary += "✅ Văn bản được soạn thảo tuân thủ quy định cơ bản."

        # 4. Lưu file DOCX (Chạy trong Thread riêng để không block server)
        filename = await asyncio.to_thread(self._save_docx, raw_content, data.contract_type)

        return {
            "content_preview": raw_content,
            "risk_report": risk_summary,
            "download_url": f"/static/docs/{filename}"
        }

    def _save_docx(self, content: str, doc_type: str) -> str:
        """
        Hàm xử lý format Word thông minh hơn
        """
        doc = Document()
        
        # Cấu hình Font mặc định
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Times New Roman'

            # --- Logic Format thông minh ---
            
            # 1. Quốc hiệu (Căn giữa, In đậm)
            if "CỘNG HÒA XÃ HỘI" in line.upper():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(13)
            
            # 2. Tiêu ngữ (Căn giữa, In đậm)
            elif "Độc lập" in line and "Tự do" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(13)
            
            # 3. Tên Hợp đồng (Căn giữa, In đậm, To) - Logic: Ngắn + Viết hoa toàn bộ
            elif len(line) < 100 and line.isupper() and "ĐIỀU" not in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(14)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)

            # 4. Các Điều khoản (Điều 1, Điều 2...) -> In đậm
            elif re.match(r"^(Điều|Khoản)\s+\d+[:\.]", line, re.IGNORECASE):
                run.bold = True
                p.paragraph_format.space_before = Pt(6)
            
            # 5. Các bên (Bên A, Bên B) -> In đậm
            elif line.startswith("Bên A") or line.startswith("Bên B") or line.startswith("ĐẠI DIỆN"):
                run.bold = True

        # Lưu file
        save_dir = "static/docs"
        os.makedirs(save_dir, exist_ok=True)
        timestamp = int(datetime.now().timestamp())
        # Clean tên file cho an toàn
        safe_type = re.sub(r'[\\/*?:"<>|]', "", doc_type).replace(" ", "_")
        filename = f"ViLaw_{safe_type}_{timestamp}.docx"
        
        doc.save(os.path.join(save_dir, filename))
        return filename